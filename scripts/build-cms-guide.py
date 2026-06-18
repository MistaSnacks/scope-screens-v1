# scripts/build-cms-guide.py
# Generates the Scope Screenings CMS editor guide as a .docx, brand-styled
# (Aachen Bold headings, Libre Franklin body) with an annotated screenshot per
# section. Screenshots come from .shots/guide/*.png (see .shots/capture-cms-guide.mjs).
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SHOTS = os.path.join(_ROOT, ".shots/guide")

RUST = RGBColor(0xB1, 0x3A, 0x2A)
INK = RGBColor(0x14, 0x12, 0x10)
GREY = RGBColor(0x55, 0x50, 0x4A)


def style_run(r, font="Libre Franklin", size=10.5, bold=False, color=INK):
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color


def heading(doc, text, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run(text), font="Aachen Bold", size=size, color=RUST)
    return p


def para(doc, text, color=INK, size=10.5, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    style_run(p.add_run(text), size=size, color=color)
    return p


def bullet(doc, text, label=None):
    p = doc.add_paragraph(style="List Bullet")
    if label:
        style_run(p.add_run(label + ": "), bold=True)
    style_run(p.add_run(text))


def legend_line(doc, n, name, desc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    style_run(p.add_run(f"{n}  "), bold=True, color=RUST, size=11)
    style_run(p.add_run(name), bold=True)
    style_run(p.add_run(" - " + desc))


def optional_image(doc, filename, caption=None, missing_note=None):
    """Embed .shots/guide/<filename> if present; otherwise leave a labelled
    placeholder so the slot is obvious until the screenshot is dropped in."""
    img = os.path.join(SHOTS, filename)
    if os.path.exists(img):
        doc.add_picture(img, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].paragraph_format.space_after = Pt(3)
        if caption:
            c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraph_format.space_after = Pt(6)
            style_run(c.add_run(caption), size=9, color=GREY)
    elif missing_note:
        n = doc.add_paragraph(); n.paragraph_format.space_after = Pt(6)
        style_run(n.add_run(missing_note), size=9, color=GREY)


# (key, heading, collections-note, [ (n, name, desc) ... ], also-note)
SECTIONS = [
    ("Hero", "Hero", "Form: Hero", [
        ("1", "Eyebrow", "small label above the title (e.g. “Feature Presentation”)."),
        ("2", "Title", "the big wordmark. Put a line break in the field for two lines."),
        ("3", "Tagline", "the line under the wordmark."),
    ], "Also in the Hero form: Poster image (the still shown before the video plays) and Video (the background reel - upload/select in the Media Manager)."),

    ("WhatIs", "What Is Scope", "Form: WhatIs", [
        ("1", "Eyebrow", "the small mono label."),
        ("2", "Title", "the section heading."),
        ("3", "Body", "the editorial paragraph."),
        ("4", "Motto", "the italic quote line."),
        ("5", "Clapperboard", "the slate lines (Production / Director / Location / Est. / Runs) are now fields right here in the WhatIs form - Clap 1–5, each a Label and a Value."),
    ], None),

    ("BuiltForAccess", "Built For Access", "Form: BuiltForAccess", [
        ("1", "Eyebrow", "the chapter label."),
        ("2", "Title", "the section heading (line break allowed)."),
        ("3", "Founder quote", "the large blockquote."),
        ("4", "Founder name", "shown on the photo frame and beside the quote."),
        ("5", "Stats", "the big numbers (200+, 150+ …) are fixed in the design - there's nothing to edit here."),
    ], "Also in the form: Founder title, Founder credential, and the Founder photo (upload)."),

    ("MagicGallery", "Scope Screenings Magic", "Form: MagicGallery  +  List: MagicGallery - Moments", [
        ("1", "Eyebrow", "the chapter label."),
        ("2", "Title", "the section heading (line break allowed)."),
        ("3", "Body", "the paragraph."),
        ("4", "Button", "the link text and where it goes (label + link)."),
    ], "The reel photos live in “MagicGallery - Moments” - one row per photo (image, badge, title, caption, order)."),

    ("Submissions", "Submissions", "Form: Submissions", [
        ("1", "Eyebrow", "the open-call label."),
        ("2", "Title", "the section heading."),
        ("3", "Intro", "the paragraph under the title."),
        ("4", "Chips", "the little meta tags are now fields in the Submissions form - Chip 1–4, each a Label and an Accent (curtain or rust)."),
        ("5", "Button + link", "the Submit button text and where it points (FilmFreeway)."),
    ], "Note: the deadline ladder and notification date are set in the code, not the CMS."),

    ("Archives", "The Archives", "Form: Archives", [
        ("1", "Eyebrow", "the chapter label."),
        ("2", "Title", "the section heading."),
        ("3", "Body", "the paragraph."),
        ("4", "Button + link", "the link text and where it goes (the schedule)."),
    ], "Note: the film thumbnails (the filmstrip) are not in the CMS yet - they live in the code for now. "
       "You can hide this whole section from Site Settings - see “Hide a section”."),

    ("Support", "Keep It Running", "Form: Support  +  List: Support - Press Kit", [
        ("1", "Title", "the section heading."),
        ("2", "Funder copy", "the “Become a Funder” card title and body."),
        ("3", "Giving tiers", "the donor chips are now fields in the Support form - Giving Tier 1–4 (a Label, plus a Featured checkbox for the highlighted one)."),
        ("4", "Donate link", "where “Support the Festival” goes - change this if your fiscal sponsor changes."),
        ("5", "Press kit", "the press rows and their download links stay in the “Support - Press Kit” list (one row per resource: label, format, link)."),
    ], "Also in the Support form: the Press card title & body, and the Press email."),

    ("Footer", "Footer", "Form: Footer", [
        ("1", "Sign-off", "the big sign-off heading."),
        ("2", "Newsletter heading", "the text above the email sign-up. (You can hide the whole sign-off + sign-up band from Site Settings - see “Hide a section”.)"),
        ("3", "Socials", "the Instagram / TikTok / YouTube links are now fields in the Footer form - Social 1–3, each a Label and a URL."),
        ("4", "Copyright", "the line at the very bottom."),
    ], "Also in the Footer form: the tagline paragraph and the contact email."),
]


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.6)
        s.left_margin = s.right_margin = Inches(0.7)

    # Cover header
    logo = os.path.join(_ROOT, "public/popcorn-logo.png")
    doc.add_picture(logo, height=Inches(0.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(h.add_run("Editing Your Website Content"), font="Aachen Bold", size=24, color=RUST)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(sub.add_run("A visual guide to the Scope Screenings CMS"), size=12, color=GREY)

    heading(doc, "The one rule")
    para(doc, "Open the CMS in your Wix dashboard → pick the collection for the part of the page you want to "
              "change → edit the fields → click Publish. Your live site updates within the hour.")

    heading(doc, "Opening the CMS")
    bullet(doc, "Sign in at wix.com and open this site's Dashboard.", label="1")
    bullet(doc, "In the left-hand menu choose CMS. (On some plans it sits under “Dev Mode” or "
                "“Developer Tools”, and it was previously called “Content Manager”.)", label="2")
    bullet(doc, "You'll land on the list of collections - one folder per part of the page. Click any "
                "collection to open it.", label="3")
    bullet(doc, "Edit the fields, then click Publish (top-right) to push it live.", label="4")
    optional_image(doc, "cms-nav.png",
                   caption="Where to find CMS in the Wix dashboard menu.",
                   missing_note="[ Screenshot to add → save as .shots/guide/cms-nav.png: the Wix dashboard "
                                "left menu with “CMS” highlighted. ]")
    optional_image(doc, "cms-collections.png",
                   caption="The collection list - one folder per section of the page.",
                   missing_note="[ Screenshot to add → save as .shots/guide/cms-collections.png: the CMS "
                                "collections list. ]")

    heading(doc, "How it's organized")
    para(doc, "Every part of the page is its own collection, named after the section. Almost everything for a section "
              "now lives in one editable card (its “form”) - including the short repeating bits like the clapperboard "
              "lines, the submission chips, the giving tiers and the social links, which are simply numbered fields "
              "(Clap 1–5, Chip 1–4, and so on). Only the things that genuinely grow over time stay as their own lists: "
              "Partners, the Magic Gallery moments, the Support press kit, and the Marquee. On each screenshot below, "
              "the numbered pins show exactly what each field controls.")

    heading(doc, "Your collections at a glance")
    para(doc, "Each item below is one collection (a “folder”) in the CMS. Most are a single editable card; "
              "the few marked “list” hold multiple rows.")
    bullet(doc, "the opening curtain - eyebrow, title, tagline, poster image, background video.", label="Hero")
    bullet(doc, "the “What Is Scope” section, including the clapperboard slate lines.", label="WhatIs")
    bullet(doc, "the founder section - quote, name, title, photo.", label="BuiltForAccess")
    bullet(doc, "the gallery heading and copy. The photos live in the “MagicGallery — Moments” list.", label="MagicGallery")
    bullet(doc, "the open-call section - intro, chips, Submit button.", label="Submissions")
    bullet(doc, "the films/archives section - heading, body, button.", label="Archives")
    bullet(doc, "the funder + press section. Press resources live in the “Support — Press Kit” list.", label="Support")
    bullet(doc, "the sign-off, newsletter heading, socials, copyright.", label="Footer")
    bullet(doc, "the partner logos (list - one row per logo).", label="Partners")
    bullet(doc, "the scrolling marquee phrases (list - one row per phrase).", label="Marquee")
    bullet(doc, "venue details, plus the Hide Archives and Hide newsletter band toggles.", label="Site Settings")
    para(doc, "“Subscribers” is created automatically by Wix to collect newsletter emails - you don't edit it by hand.",
         color=GREY, size=9.5)

    for key, title, collections, items, also in SECTIONS:
        doc.add_page_break()
        heading(doc, title)
        c = doc.add_paragraph(); c.paragraph_format.space_after = Pt(6)
        style_run(c.add_run(collections), size=9.5, bold=True, color=GREY)
        img = os.path.join(SHOTS, f"{key}.png")
        if os.path.exists(img):
            doc.add_picture(img, width=Inches(7.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.paragraphs[-1].paragraph_format.space_after = Pt(6)
        for n, name, desc in items:
            legend_line(doc, n, name, desc)
        if also:
            a = doc.add_paragraph(); a.paragraph_format.space_before = Pt(4)
            style_run(a.add_run(also), size=10, color=GREY)

    doc.add_page_break()
    heading(doc, "Hide a section")
    para(doc, "Two parts of the page can be hidden from the Site Settings collection - no design help "
              "needed. Open Site Settings and you'll see two checkboxes:")
    bullet(doc, "check this to hide “The Archives” block on the homepage.", label="Hide Archives section")
    bullet(doc, "check this to hide the whole footer sign-off band - the “See You At The Movies” heading and the email sign-up together.", label="Hide newsletter band")
    para(doc, "Both are unchecked by default, so both sections show. Check one and click Publish - it "
              "disappears within the hour; uncheck it to bring it back. (These two sections aren't fully "
              "built out yet, so this lets you keep them hidden until you're ready to use them.)")

    doc.add_page_break()
    heading(doc, "Photos & video")
    bullet(doc, "In any Image or Video field, upload or select in the Wix Media Manager. Leave it empty and the "
                "built-in default shows - nothing breaks.")
    bullet(doc, "Partner logos work the same way: upload in the Partners list, or leave empty to use the built-in logo.")

    heading(doc, "Leave these alone (they update themselves)")
    bullet(doc, "Tickets and the Schedule come straight from Wix Events - manage those where you already sell tickets.")
    bullet(doc, "Submission deadlines and the notification date are set in the code.")
    bullet(doc, "The marquee's “Now Showing” date is automatic from your next event.")

    heading(doc, "Good to know")
    bullet(doc, "Every list row has an Order number - lower numbers show first.")
    bullet(doc, "Two lines in one field (e.g. a title): press Enter/Shift+Enter inside the field to add the line break.")
    bullet(doc, "Don't see your change? Give it up to an hour, or republish.")

    out = os.path.join(_ROOT, "docs/Scope-Screenings-CMS-Guide.docx")
    doc.save(out)
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
